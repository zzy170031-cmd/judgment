from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Judgment Skill 功能完整汇报.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
GREEN = "1F7A4D"
RED = "9B1C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D9E2EF", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_format(paragraph, before=0, after=6, line_spacing=1.10) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    set_paragraph_format(p)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    set_paragraph_format(p, after=8, line_spacing=1.167)


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_format(p, after=4)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    p2 = cell.add_paragraph(body)
    set_paragraph_format(p2, after=0)
    doc.add_paragraph()


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        set_paragraph_format(p, after=0)
        r = p.add_run(header)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(INK)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            set_paragraph_format(p, after=0)
            p.add_run(value)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Judgment Skill 功能完整汇报")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("777777")


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Judgment Skill 功能完整汇报")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("当前定位、核心能力、项目运行拓扑、证据闭环与阶段性升级边界")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("555555")

    meta = doc.add_paragraph()
    set_paragraph_format(meta, after=12)
    meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}    适用环境：Codex 内使用    来源仓库：judgment")

    add_callout(
        doc,
        "核心结论",
        "Judgment 目前的主要作用是 Codex 项目工作流总控治理：负责路由、规划、拆分、Agent 编组、证据闭环、Git/worktree 安全推进、555 审查触发和 HTML 项目运行可视化。外部模型/工具经验吸收只是当前升级阶段的方法，不是日常核心运行功能。",
    )

    add_heading(doc, "1. 当前定位", 1)
    add_body(doc, "Judgment 是一个 Codex-only 的项目工作流治理技能包。它不追求成为多模型运行平台，也不引入外部工具的完整 runtime；它把项目推进过程中需要的规划、拆工、审查、证据、Git 安全和可视化运行状态收束成一套可复用的 Codex 技能与规则。")
    add_body(doc, "它的核心价值不是替代具体实现技能，而是在开始做重活之前判断应该走哪条路线、启用哪些技能、需要什么证据、哪些动作必须被阻止或升级审查。")

    add_heading(doc, "2. 核心能力总览", 1)
    add_matrix_table(
        doc,
        ["能力", "主要用途", "典型输出"],
        [
            ["666 总控路由", "判断任务应该直接回答、规划、拆工、审查、分线程、自动化还是升级到 555。", "路由结论、Gate、启用技能、禁止事项、下一步最小 gate。"],
            ["work-planner", "把模糊或中大型目标整理成 Codex 可执行计划。", "需求澄清、XA/XB Gate、执行路线、证据要求、拆工建议。"],
            ["needs-solution-designer", "把不清晰的业务想法、产品需求或流程痛点澄清为稳定方案。", "确认事实、工作假设、成功标准、复用/新建判断。"],
            ["work-splitter", "把清晰的大任务拆成泳道、Agent 小组、子任务契约和验收路径。", "工作泳道、Agent 编组、输入输出、验证命令、handoff。"],
            ["555 五代理闭环", "对高风险 done claim、release、架构、安全、后端/shared surface 做证据审查。", "五席审查、Core Challenger 质疑、Audit Specialist 证据、go/no-go。"],
            ["项目拓扑与 HTML cockpit", "把真实或计划中的 Agent、节点、测试、证据、blocker 可视化。", "project-agent-graph JSON、HTML 泳道流式执行页面。"],
            ["Git / worktree 安全", "保护 dirty worktree、并行 worker、branch ownership、集成冲突和提交边界。", "worktree 分配、dirty ownership、集成 owner、禁止 Git 操作。"],
            ["Browser / UI 验证", "对可视化、HTML、本地 Web、交互和响应式布局要求可见证据。", "Browser 预览、截图/DOM 证据、条件通过或阻塞说明。"],
        ],
        [1700, 4300, 3360],
    )

    add_heading(doc, "3. 666 总控路由", 1)
    add_body(doc, "666 是 Judgment 的上游路由器。它负责在任务变重之前先判断当前应该用最小的有效路径推进。")
    for item in [
        "判断任务是直接回答、只读分析、进入规划、进入拆工、进入实现，还是升级审查。",
        "识别当前 XA/XB 开发线和 G0-G6 交付 Gate。",
        "判断是否需要证据 ledger、Browser 验证、安全审查、Git worktree 或上下文 handoff。",
        "判断是否触发 555、Core Challenger、Audit Specialist 或 worker 分派。",
        "避免小任务被包装成重流程，也避免高风险任务被轻率执行。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. 规划与需求澄清", 1)
    add_body(doc, "work-planner 和 needs-solution-designer 共同处理“还不该直接开工”的阶段。前者偏完整项目推进计划，后者偏需求本身尚不清晰时的方案澄清。")
    for item in [
        "把用户目标转成明确的完成标准和非目标。",
        "分离已确认事实、工作假设、风险和待决策项。",
        "决定复用现有技能、轻改现有技能、新建技能/Agent，还是保持简单流程。",
        "在进入实现前定义验收方式、证据路径和下一步最小 gate。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. 工作拆分与 Agent 编组", 1)
    add_body(doc, "work-splitter 负责把已经相对清晰的大任务拆成可执行泳道。拆分标准不是岗位名称本身，而是责任边界、可验证输出、文件/系统写入范围和下游接收者。")
    add_matrix_table(
        doc,
        ["泳道 / Agent 类型", "创建条件", "验证方式"],
        [
            ["Controller", "任何项目拓扑存在时。", "图谱完整性、Gate、blocker、集成状态。"],
            ["Product / Spec", "目标、范围、用户流、验收标准不清晰。", "产品验收清单和非目标。"],
            ["Frontend / UX", "涉及 UI、交互、响应式、可访问性或设计还原。", "Browser 流程、截图、客户端测试。"],
            ["Backend / Fullstack", "涉及 API、数据、权限、持久化或纵向切片。", "后端测试、契约测试、Browser + API 证据。"],
            ["QA / Security", "涉及回归、发布、安全、外部发送、生产或 AI 工具边界。", "测试报告、威胁模型、缓解证据。"],
            ["Git / Integration", "多个 worker、worktree、共享文件或合并冲突存在。", "Git 状态、diff、CI、集成清单。"],
            ["Release / Ops", "涉及发布、监控、回滚、支持或事故。", "发布清单、监控、回滚证据。"],
        ],
        [2200, 3900, 3260],
    )

    add_heading(doc, "6. 项目运行拓扑与 HTML Cockpit", 1)
    add_body(doc, "当前 Judgment 已具备项目拓扑生成与 HTML 可视化能力。它可以把一个项目表示成节点、泳道、Agent、测试、证据、blocker、handoff、release decision 和 Git/integration 风险。")
    for item in [
        "每个工作节点必须有上游输入、下游消费者和测试/证据。",
        "每个实现节点必须连接测试或证据节点，不能凭信心标记 done。",
        "真实子代理必须记录 agent_id、delegation_tool、conversation_surface、last_message 和 next_input。",
        "HTML cockpit 是状态镜像，不是证据本身；证据仍来自测试、Git、日志、截图、侧聊和审查结论。",
        "当前渲染器支持浅色 Agent 办公室风格、泳道流式执行、右侧 Agent 状态、最近任务和活动动态。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. 证据、测试与上下游咬合", 1)
    add_body(doc, "Judgment 的一个关键原则是：每一步都必须能被测试或证据承接，并且要和前后节点互相咬合。")
    for item in [
        "没有 test_or_evidence 的节点不能进入 done。",
        "测试节点必须有 pass/fail oracle 和 owner。",
        "handoff 必须有接收者和预期输出。",
        "blocked 节点必须说明 blocked_by、blocking_reason 或 required_human_input。",
        "重大 done、release、backend/shared surface、安全或 AI/Agent 结论必须有独立审查或 555。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Git、Worktree 与 Dirty Ownership", 1)
    add_body(doc, "当任务涉及 Git、并行 worker 或不干净工作区时，Judgment 会先保护真实文件状态。")
    for item in [
        "创建或分配 worker 前先确认 branch、HEAD、dirty state 和 worktree ownership。",
        "一个写入 worker 对应一个分支和一个路径，避免共享写入范围。",
        "branch 已被其他 worktree checkout 时视为 ownership 冲突，而不是普通错误。",
        "提交、清理、reset、remove、prune、push 前必须判断 dirty 文件属于当前任务、用户、生成物还是未知。",
        "未知 ownership 默认保留，不做破坏性操作。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. Browser、安全、工具组合与上下文保护", 1)
    add_body(doc, "Judgment 会根据任务性质决定是否需要 Browser 可见验证、安全审查、工具组合选择或上下文 handoff。")
    for item in [
        "UI、HTML、本地 Web、交互、设计还原和响应式布局需要 Browser 证据。",
        "auth、权限、支付、用户数据、secrets、外部发送、生产、破坏性动作和 AI 工具边界需要安全 gate。",
        "选择 plugin、skill、MCP/connector、script、automation 或 no-install 时走 tool portfolio 判断。",
        "上下文压力高时先输出 checkpoint/handoff，再继续长任务。",
        "重复、高成本、稳定输入的流程才考虑封装为 skill、automation、subagent、script 或 plugin。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. 阶段性升级机制：外部优秀经验吸收", 1)
    add_callout(
        doc,
        "定位修正",
        "外部模型/工具经验吸收不是 Judgment 的核心运行功能。它只是当前升级阶段的方法，用来参考 ChatGPT、Codex、Claude Code、Gemini、GLM、Git/GitHub、字节系工具等优秀实践，并把适合的部分转译成 Judgment 内部规则、技能、脚本或可视化机制。",
        "FFF4DF",
    )
    for item in [
        "只吸收可提升规划、证据、拆工、审查、可视化、Git 安全和 Agent 编组的稳定模式。",
        "不直接复制其他工具的命令体系、模型假设、隐藏推理存储、runtime daemon 或组织流程。",
        "所有落地仍然以 Codex 内可执行 surface 为准。",
        "升级完成后，这一机制应作为维护和演进方法，而不是日常项目推进主路径。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. 当前仓库资产", 1)
    add_matrix_table(
        doc,
        ["资产", "作用"],
        [
            ["SKILL.md / skills/666", "兼容入口和总控路由器。"],
            ["skills/work-planner", "完整项目规划入口。"],
            ["skills/needs-solution-designer", "需求澄清与方案蓝图。"],
            ["skills/work-splitter", "工作拆分、Agent 编组和线程策略。"],
            ["skills/555", "五代理闭环审查和证据验证。"],
            ["rules/project-agent-topology-standard.md", "项目拓扑、Agent 状态、HTML cockpit 和证据咬合规则。"],
            ["rules/git-worktree-standard.md", "worktree 隔离、dirty ownership 和 Git 生命周期安全。"],
            ["rules/browser-flow-testing-standard.md", "Web/UI/HTML 可见验证规则。"],
            ["rules/security-review-standard.md", "权限、数据、生产、外部发送和 AI 工具安全规则。"],
            ["scripts/render-project-agent-graph.py", "把 graph JSON 渲染为 HTML Agent 办公室 / 泳道流式 cockpit。"],
            ["templates/project-agent-graph.example.json", "项目拓扑示例数据。"],
        ],
        [3200, 6160],
    )

    add_heading(doc, "12. 明确边界", 1)
    add_body(doc, "Judgment 当前不应该被理解为完整外部 runtime、多模型执行平台、自动绕过 sandbox 的系统、隐藏状态机或假装 live chat 的 HTML 页面。")
    for item in [
        "它不自动创建 N 个 Agent 充图表；Agent 必须来自真实项目需求。",
        "它不把 HTML 当证据；HTML 只是镜像和导航。",
        "它不允许实现 Agent 自己批准 release、backend/shared surface、安全或 AI/Agent 风险结论。",
        "它不盲目安装外部工具或复制外部模型功能。",
        "它的核心仍然是 Codex 内的项目治理、证据闭环和安全推进。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "13. 当前可用结论", 1)
    add_body(doc, "当前 Judgment 已经从单一工作流技能进化为 Codex 项目治理包。它能帮助 Codex 在项目推进时先判断路线，再拆分任务，再绑定证据，再决定是否启用 Agent、Browser、Git worktree、555 或自动化。它最重要的价值是让项目推进更稳、更可验证、更不容易在上下文、Git 状态、审查证据和多线协作中失控。")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
